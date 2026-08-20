#!/usr/bin/env python3
"""Generate partner-facing Liquid AI NLI evaluation memo (Word)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "E8E8E8")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, text in enumerate(row):
            cells[c_idx].text = text
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_labeled_paragraph(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label} ")
    r.bold = True
    p.add_run(value)


def main() -> None:
    out_dir = Path(__file__).resolve().parent
    out = out_dir / "Liquid-AI-LFM25-Encoder-NLI-Eval-Memo-2026-08-20.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(
        "LFM2.5-Encoder NLI evaluation — results & production recommendation",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for label, value in [
        ("To:", "Liquid AI — Partnership / Applied Research"),
        ("From:", "Jean-Baptiste Dézard"),
        (
            "Re:",
            "LFM2.5-Encoder as NLI CrossEncoder replacement in SGRS (governed multi-agent stack)",
        ),
        ("Date:", "20 August 2026"),
    ]:
        add_labeled_paragraph(doc, label, value)

    doc.add_paragraph()

    doc.add_heading("Executive summary", level=1)
    doc.add_paragraph(
        "We completed an end-to-end evaluation of LFM2.5-Encoder NLI against our production "
        "CrossEncoder baseline (cross-encoder/nli-deberta-v3-large) on a frozen 57-pair "
        "governance gold harness and an expanded 1,006-pair domain corpus (711 train / 295 eval; "
        "FR/EN M&A, insurance, clinical, AML, cyber)."
    )
    p = doc.add_paragraph()
    r = p.add_run("Recommendation: ")
    r.bold = True
    p.add_run(
        "Go on LFM2.5-Encoder-230M with our v3 training recipe (HF L4 checkpoint, ~5 min train). "
        "The model passes all safety gates, improves governance routing accuracy by +15.8 pp vs "
        "DeBERTa v3 large (82.5% vs 66.7%), and is integrated into our facts-worker as an opt-in "
        "backend (NLI_BACKEND=liquidai). When NLI flags a contradiction, we block silent auto-merge "
        "and route the pair into human-in-the-loop (HITL) resolution — not a reject-and-forget gate."
    )
    doc.add_paragraph(
        "350M on the same corpus beats 230M on routing (87.7% vs 82.5%) with 100% block recall on "
        "our best Mac MPS run; it is a credible optional upsize (~1.5× parameters). L4 training for "
        "350M requires recipe alignment with MPS (refine weight decay 0.1, lower domain LR) — default "
        "L4 hyperparameters over-fit the refine stage and hurt gray-zone HITL routing."
    )
    doc.add_paragraph(
        "Embeddings / pgvector replacement was explicitly out of scope for this cycle."
    )

    doc.add_heading("Context — where the encoder sits in SGRS", level=1)
    doc.add_paragraph(
        "SGRS (Swarm of Governed Agents) is a research architecture for long-horizon, auditable "
        "multi-agent knowledge work (due diligence, compliance, pharmacovigilance). LLMs propose "
        "claims; a deterministic governance kernel decides equivalence, contradiction handling, and "
        "finality. The encoder NLI model is Couche 2 in our equivalence stack:"
    )
    add_bullets(
        doc,
        [
            "Couche 0 — typed canonicalisation (deterministic)",
            "Couche 1 — embeddings (gray-zone pre-filter only; unchanged in this eval)",
            "Couche 2 — NLI via facts-worker /nli (this evaluation target)",
            "Couche 3 — governance decision (deterministic approve/deny on merge proposals)",
        ],
    )
    doc.add_paragraph(
        "The expensive failure mode is a false merge (silently merging incompatible facts). NLI's job "
        "is precision on that gate. Contradictions block auto-merge and surface in the semantic graph "
        "for human resolution."
    )

    doc.add_heading("Governance routing — block routes to HITL", level=1)
    doc.add_paragraph(
        "NLI is a safety gate on auto-merge, not the final arbiter. A block means do not silently "
        "merge — it does not mean drop the pair."
    )
    add_table(
        doc,
        ["Outcome", "Auto-merge?", "Next step in SGRS"],
        [
            ["Equivalent (confidence ≥ 0.77)", "Yes", "Governed equivalent_to edge"],
            [
                "Equivalent (low confidence), neutral, accrual / refinement",
                "No",
                "HITL — human adjudicates paraphrase or scope change",
            ],
            [
                "Contradiction / refutation",
                "No",
                "Block auto-merge → contradicts edges in semantic graph → HITL/MITL "
                "(keep prior, accept update, defer, or add resolves edge)",
            ],
            ["Unrelated (false-positive trap)", "No", "No merge; no contradiction workflow"],
            ["NLI unavailable", "No", "Fail-closed neutral — never merge on guess"],
        ],
    )
    add_bullets(
        doc,
        [
            "False-merge (hard gate 0%): incompatible facts merged without human review.",
            "Block recall: NLI flags contradiction → blocks auto-merge so humans can resolve.",
            "HITL routing: gray-zone pairs reach human review without misclassification as hard contradiction.",
            "Accrual over-block: failure mode — accrual/refinement treated as contradiction instead of HITL.",
        ],
    )

    doc.add_heading("What we measured", level=1)
    add_table(
        doc,
        ["Layer", "Description"],
        [
            [
                "Gold harness",
                "57 frozen pairs; minConfidence 0.77; metrics: false-merge, block recall, "
                "HITL routing, accrual over-block",
            ],
            ["B4 gate", "Macro-F1 on held-out eval.jsonl (n=295) vs DeBERTa baseline 0.730"],
            [
                "Training pipeline",
                "3 stages: MNLI probe → domain fine-tune (wd=0.1) → hard-negative refine (wd=0.01)",
            ],
            [
                "Integration",
                "facts-worker /nli; bidirectional mutual entailment; same JSON contract as DeBERTa; "
                "fail-closed on load/inference failure",
            ],
        ],
    )

    doc.add_heading("Domain training corpus", level=1)
    add_table(
        doc,
        ["Stat", "Value"],
        [
            ["Total pairs", "1,006"],
            ["Train / eval", "711 / 295"],
            ["Labels", "~338 each (equivalent · neutral · contradiction)"],
            ["Languages", "EN 580 · FR 415 · DE 11"],
            ["Sources", "Hand-curated seeds + synthetic domain generator + gold-failure mining"],
        ],
    )
    doc.add_paragraph(
        "Hugging Face dataset: huggingface.co/datasets/jeanbaptdzd/liquid-nli-domain-1k"
    )

    doc.add_heading("Gold harness results (n=57, minConf 0.77)", level=1)
    doc.add_paragraph(
        "Primary comparison: governance routing accuracy and safety gates vs DeBERTa v3 large."
    )
    add_table(
        doc,
        [
            "Model",
            "Hub checkpoint",
            "Routing",
            "False-merge",
            "Block recall",
            "HITL routing",
            "Accrual over-block",
        ],
        [
            ["DeBERTa v3 large", "production baseline", "66.7%", "0.0%", "94.1%", "50.0%", "—"],
            [
                "LFM 230M v3 (L4)",
                "jeanbaptdzd/lfm25-nli-v3-calibrated-l4",
                "82.5%",
                "0.0%",
                "94.1%",
                "85.7%",
                "15.4%",
            ],
            [
                "LFM 350M (MPS, 1k)",
                "local (not on Hub)",
                "87.7%",
                "0.0%",
                "100%",
                "100%",
                "0.0%",
            ],
            [
                "LFM 350M (L4, MPS-recipe retune)",
                "jeanbaptdzd/lfm25-nli-350m-v1k-mps-recipe-l4",
                "86.0%",
                "0.0%",
                "100%",
                "92.9%",
                "7.7%",
            ],
            [
                "LFM 350M (L4, default recipe)",
                "jeanbaptdzd/lfm25-nli-350m-v1k-calibrated-l4",
                "77.2%",
                "0.0%",
                "94.1%",
                "64.3%",
                "38.5%",
            ],
        ],
    )
    doc.add_paragraph(
        "Safety gates: false-merge = 0% (hard); block recall ≥ 94.1% (match DeBERTa). "
        "Production pick: 230M v3 L4 — best balance of safety, routing gain, CPU footprint, "
        "and reproducible Hub artifact."
    )

    doc.add_heading("Held-out eval macro-F1 (n=295)", level=1)
    add_table(
        doc,
        ["Model", "macro-F1", "vs DeBERTa (0.730)"],
        [
            ["DeBERTa v3 large", "0.730", "baseline"],
            ["LFM 230M v3", "0.959", "+0.229"],
            ["LFM 350M (MPS / aligned L4)", "~0.959–0.966", "+0.229–0.236"],
        ],
    )
    doc.add_paragraph("Gate: challenger ≥ baseline — PASS for all fine-tuned LFM runs.")

    doc.add_heading("Training runs summary", level=1)
    add_table(
        doc,
        ["Run", "Encoder", "Infra", "Corpus", "Hub output", "Verdict"],
        [
            ["v3", "230M", "HF L4 ~5 min", "1k", "lfm25-nli-v3-calibrated-l4", "Go — production candidate"],
            ["350M", "350M", "Mac MPS ~14 min", "1k", "local only", "Go — optional upsize"],
            ["350M", "350M", "HF L4 default", "1k", "lfm25-nli-350m-v1k-calibrated-l4", "No-go (HITL over-block)"],
            ["350M", "350M", "HF L4 MPS-recipe", "1k", "lfm25-nli-350m-v1k-mps-recipe-l4", "Near MPS gold scores"],
        ],
    )
    doc.add_paragraph(
        "Recipe lesson (350M on L4): Stage-3 refine with wd=0.01 and 2× domain learning rate drove "
        "refine train_loss to 0.001 (vs 0.06 on MPS), collapsing gray-zone pairs to high-confidence "
        "contradiction. Retuning with refine wd=0.1, domain lr=1e-5, and refine-only from the MPS "
        "domain checkpoint restored gold performance."
    )
    doc.add_paragraph(
        "Rejected paths: 350M on small (~186-pair) corpus (block recall 76.5%); aggressive Stage-3 "
        "refine on 230M (block recall 82.4%); zero-shot LFM probe (gates fail as expected)."
    )

    doc.add_heading("Published artifacts (Hugging Face)", level=1)
    add_bullets(
        doc,
        [
            "Production checkpoint (230M v3): huggingface.co/jeanbaptdzd/lfm25-nli-v3-calibrated-l4",
            "Domain corpus (1k pairs): huggingface.co/datasets/jeanbaptdzd/liquid-nli-domain-1k",
            "Training scripts: huggingface.co/jeanbaptdzd/liquid-nli-scripts",
            "350M L4 retune: huggingface.co/jeanbaptdzd/lfm25-nli-350m-v1k-mps-recipe-l4",
        ],
    )

    doc.add_heading("Integration status", level=1)
    doc.add_paragraph(
        "230M v3 is wired into facts-worker for opt-in staging trials. Default backend remains "
        "DeBERTa CrossEncoder. Smoke tests (paraphrase + contradiction pairs via live /nli) pass."
    )
    add_bullets(
        doc,
        [
            "NLI_BACKEND=liquidai",
            "LIQUID_NLI_MODE=finetuned",
            "EQUIV_MIN_CONFIDENCE=0.77",
            "Base model: LiquidAI/LFM2.5-Encoder-230M + fine-tuned 3-class head",
        ],
    )

    doc.add_heading("Per-component verdict", level=1)
    add_table(
        doc,
        ["Component", "Verdict", "Notes"],
        [
            [
                "LFM2.5-Encoder-230M NLI (v3)",
                "Go",
                "Safety-equivalent to DeBERTa; +15.8 pp routing; Hub checkpoint; integrated",
            ],
            [
                "LFM2.5-Encoder-350M NLI",
                "Go with caveats",
                "Best MPS gold scores; L4 needs MPS-aligned recipe",
            ],
            ["LFM embeddings", "Deferred", "No pgvector change in this cycle"],
            ["Flip production default", "Separate step", "Env-flag trial on staging recommended first"],
        ],
    )

    doc.add_heading("Suggested next steps", level=1)
    add_numbered(
        doc,
        [
            "Joint review of the v3 checkpoint and remaining gold routing misses (~10 pairs, mostly "
            "typed-dimension paraphrases where Couche 0 should merge before NLI).",
            "Liquid-side review of 350M training recipe (refine weight decay, probe size, SNLI mix) "
            "toward a single canonical L4 config.",
            "CPU latency benchmark — P50/P95 /nli vs DeBERTa on identical hardware (230M already "
            "qualitatively faster in our runs).",
            "Controlled staging rollout; compare HITL rate and finality outcomes vs CrossEncoder.",
        ],
    )

    doc.add_heading("Ask", level=1)
    doc.add_paragraph(
        "We have a reproducible, safety-certified, Hub-published LFM2.5-Encoder-230M NLI checkpoint "
        "that improves governance routing over our DeBERTa production baseline without false merges. "
        "We would value Liquid AI feedback on the training recipe and your interest in a deeper "
        "350M / multilingual iteration aligned with your encoder roadmap."
    )

    doc.add_heading("Try SGRS (invitation)", level=1)
    doc.add_paragraph(
        "The governed swarm stack is open source. You can run the full multi-agent demo "
        "(contradiction detection, governance gates, HITL finality review) without any Liquid-specific "
        "integration — it uses the same architecture this NLI eval exercised."
    )
    add_labeled_paragraph(
        doc,
        "Public repository:",
        "github.com/DealExMachina/open-governed-swarm-of-agents",
    )
    add_labeled_paragraph(
        doc,
        "Liquid NLI integration branch:",
        "on our private dev branch today (not yet merged to the public repo). We can publish a "
        "public feature branch or grant read access if you want to replay the NLI gold harness and "
        "facts-worker /nli swap end-to-end.",
    )
    doc.add_paragraph("Quick start on the public repo (DeBERTa NLI or NLI-off defaults):")
    add_bullets(
        doc,
        [
            "git clone https://github.com/DealExMachina/open-governed-swarm-of-agents.git",
            "cp .env.example .env — set OPENAI_API_KEY or Ollama",
            "docker compose up -d && pnpm install",
            "CHECK_SERVICES_MAX_WAIT_SEC=300 pnpm run check:services",
            "pnpm run ensure-schema && pnpm run ensure-bucket && pnpm run ensure-stream",
            "pnpm run seed:all && pnpm run bootstrap-once",
            "pnpm run swarm:start — then pnpm run demo (Project Horizon M&A scenario)",
        ],
    )
    doc.add_paragraph(
        "Demo walkthrough: demo/DEMO.md — covers contradiction edges, governance blocks, and "
        "human-in-the-loop finality. NLI equivalence gate is opt-in (EQUIVALENCE_GATE=1, SKIP_NLI=0)."
    )
    doc.add_paragraph(
        "Liquid NLI trial path (once branch is public or shared): set NLI_BACKEND=liquidai, "
        "LIQUID_NLI_MODE=finetuned, pull jeanbaptdzd/lfm25-nli-v3-calibrated-l4 into "
        "workers/facts-worker/checkpoints/nli-domain-v3-calibrated, then run "
        "scripts/run-liquid-nli-gold-eval.sh or scripts/run-liquid-nli-issue06-smoke.sh."
    )

    doc.add_paragraph()
    p = doc.add_paragraph("Jean-Baptiste Dézard")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(out)
    print(out)

    # Keep legacy filename in sync for internal references
    legacy = out_dir / "LIQUIDAI-NLI-EVAL-WRAPUP-2026-08-20.docx"
    if legacy != out:
        doc.save(legacy)
        print(legacy)


if __name__ == "__main__":
    main()
